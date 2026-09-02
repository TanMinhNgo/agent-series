"""Document ingestion and pgvector retrieval for the local knowledge base."""

from __future__ import annotations

import hashlib
import re
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import select

from .storage import Database, Document, DocumentChunk, KnowledgeCollection, KnowledgeCollectionDocument, current_user_id, document_scope_key
from .file_storage import FileStorageService
from .tools.base import ToolSpec

CHUNK_SIZE = 800
CHUNK_OVERLAP = 120
DOCX_SUFFIX = ".docx"
ALLOWED_DOCUMENT_SUFFIXES = {".pdf", DOCX_SUFFIX, ".md"}
NO_DOCUMENTS_RESULT = "Không có tài liệu nào đã index để trả lời câu hỏi này."


def _chunks(text: str) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )
    return [chunk.strip() for chunk in splitter.split_text(text) if chunk.strip()]


def extract_document_parts(path: Path, suffix: str) -> tuple[list[tuple[int, str]], int]:
    """Return indexable text with a real page number or a virtual section number."""
    suffix = suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return (
            [
                (page_number, chunk)
                for page_number, page in enumerate(reader.pages, start=1)
                for chunk in _chunks(page.extract_text() or "")
            ],
            len(reader.pages),
        )
    if suffix == DOCX_SUFFIX:
        from docx import Document as DocxDocument

        text = "\n".join(item.text for item in DocxDocument(BytesIO(path.read_bytes())).paragraphs)
    elif suffix == ".md":
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError("Chỉ hỗ trợ file PDF, DOCX hoặc Markdown.")
    chunks = _chunks(text)
    return list(enumerate(chunks, start=1)), len(chunks)


def extract_document_parts_bytes(data: bytes, suffix: str) -> tuple[list[tuple[int, str]], int]:
    """Extract document text from private remote storage without a temp public URL."""
    suffix = suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(data))
        return ([(page_number, chunk) for page_number, page in enumerate(reader.pages, start=1) for chunk in _chunks(page.extract_text() or "")], len(reader.pages))
    if suffix == DOCX_SUFFIX:
        from docx import Document as DocxDocument
        text = "\n".join(item.text for item in DocxDocument(BytesIO(data)).paragraphs)
    elif suffix == ".md":
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError("Chỉ hỗ trợ file PDF, DOCX hoặc Markdown.")
    chunks = _chunks(text)
    return list(enumerate(chunks, start=1)), len(chunks)


class KnowledgeService:
    def __init__(self, database: Database, knowledge_dir: Path, embedding_model: str, storage: FileStorageService | None = None):
        self.database = database
        self.knowledge_dir = knowledge_dir
        self.embedding_model_name = embedding_model
        self._embedder = None
        self.storage = storage or FileStorageService(knowledge_dir)

    def _embedder_instance(self):
        if self._embedder is None:
            from sentence_transformers import SentenceTransformer
            self._embedder = SentenceTransformer(self.embedding_model_name)
        return self._embedder

    def _embed(self, values: list[str], prefix: str) -> list[list[float]]:
        vectors = self._embedder_instance().encode(
            [f"{prefix}: {value}" for value in values], normalize_embeddings=True
        )
        return vectors.tolist()

    def list_documents(self, project_id: str | None = None) -> list[Document]:
        with self.database.session() as session:
            statement = select(Document).order_by(Document.created_at.desc())
            if project_id is not None:
                statement = statement.where(Document.project_id == project_id)
            return list(session.scalars(statement))

    def ensure_remote(self, document_id: str) -> Document | None:
        with self.database.session() as session:
            document = session.get(Document, document_id)
            if document is None or document.storage_provider != "local":
                return document
            stored = self.storage.migrate_local(document.stored_name, document.original_name, "knowledge")
            if stored is None:
                return document
            document.storage_provider, document.stored_name, document.storage_file_id = stored.provider, stored.stored_name, stored.file_id
            session.commit()
            return document

    def list_collections(self, project_id: str) -> list[KnowledgeCollection]:
        with self.database.session() as session:
            return list(session.scalars(select(KnowledgeCollection).where(KnowledgeCollection.project_id == project_id).order_by(KnowledgeCollection.updated_at.desc())))

    def get_collection(self, collection_id: str) -> KnowledgeCollection | None:
        with self.database.session() as session:
            return session.get(KnowledgeCollection, collection_id)

    def create_collection(self, project_id: str, name: str, description: str | None = None) -> KnowledgeCollection:
        clean_name = name.strip()
        if not clean_name:
            raise ValueError("Tên collection không được để trống.")
        with self.database.session() as session:
            item = KnowledgeCollection(project_id=project_id, name=clean_name[:160], description=(description or "").strip()[:10_000] or None)
            session.add(item); session.commit(); return item

    def update_collection(self, collection_id: str, name: str, description: str | None = None) -> KnowledgeCollection | None:
        clean_name = name.strip()
        if not clean_name:
            raise ValueError("Tên collection không được để trống.")
        with self.database.session() as session:
            item = session.get(KnowledgeCollection, collection_id)
            if item is None: return None
            item.name, item.description = clean_name[:160], (description or "").strip()[:10_000] or None
            session.commit(); return item

    def delete_collection(self, collection_id: str) -> bool:
        with self.database.session() as session:
            item = session.get(KnowledgeCollection, collection_id)
            if item is None: return False
            session.delete(item); session.commit(); return True

    def collection_documents(self, collection_id: str) -> list[Document]:
        with self.database.session() as session:
            return list(session.scalars(select(Document).join(KnowledgeCollectionDocument).where(KnowledgeCollectionDocument.collection_id == collection_id).order_by(Document.original_name)))

    def set_collection_documents(self, collection_id: str, document_ids: list[str]) -> list[Document]:
        with self.database.session() as session:
            collection = session.get(KnowledgeCollection, collection_id)
            if collection is None: raise ValueError("Không tìm thấy collection.")
            documents = list(session.scalars(select(Document).where(Document.id.in_(set(document_ids)))).all()) if document_ids else []
            if len(documents) != len(set(document_ids)) or any(item.project_id != collection.project_id for item in documents):
                raise ValueError("Tài liệu phải thuộc cùng Project với collection.")
            session.query(KnowledgeCollectionDocument).filter(KnowledgeCollectionDocument.collection_id == collection_id).delete()
            session.add_all(KnowledgeCollectionDocument(collection_id=collection_id, document_id=item.id) for item in documents)
            session.commit(); return documents

    def upload(self, original_name: str, data: bytes, project_id: str | None = None) -> tuple[Document, bool]:
        suffix = Path(original_name).suffix.lower()
        if suffix not in ALLOWED_DOCUMENT_SUFFIXES:
            raise ValueError("Chỉ nhận file PDF, DOCX hoặc Markdown (.md).")
        if not data:
            raise ValueError("File đang trống.")
        if len(data) > 25 * 1024 * 1024:
            raise ValueError("File vượt giới hạn 25 MB.")
        digest = hashlib.sha256(data).hexdigest()
        scope_key = document_scope_key(project_id)
        self.knowledge_dir.mkdir(parents=True, exist_ok=True)
        with self.database.session() as session:
            existing = session.scalar(select(Document).where(Document.sha256 == digest, Document.scope_key == scope_key))
            if existing:
                return existing, False
            stored = self.storage.upload(data, original_name, "knowledge")
            document = Document(
                original_name=Path(original_name).name,
                stored_name=stored.stored_name,
                storage_provider=stored.provider,
                storage_file_id=stored.file_id,
                sha256=digest,
                scope_key=scope_key,
                project_id=project_id,
            )
            session.add(document)
            session.commit()
            return document, True

    def index(self, document_id: str) -> Document:
        with self.database.session() as session:
            document = session.get(Document, document_id)
            if document is None:
                raise ValueError("Không tìm thấy tài liệu.")
            document.status, document.error = "indexing", None
            session.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
            session.commit()
            try:
                suffix = Path(document.original_name).suffix.lower()
                parts, unit_count = extract_document_parts_bytes(
                    self.storage.read(document.storage_provider, document.stored_name, document.storage_file_id), suffix
                )
                if not parts:
                    if suffix == ".pdf":
                        document.status, document.error = "needs_ocr", "PDF không có text layer. Hãy OCR trước rồi thử index lại."
                    else:
                        document.status, document.error = "failed", "Tài liệu không có nội dung văn bản để index."
                    session.commit()
                    if document.status == "failed":
                        raise RuntimeError(document.error)
                    return document
                embeddings = self._embed([content for _, content in parts], "passage")
                for index, ((page_number, content), embedding) in enumerate(zip(parts, embeddings)):
                    session.add(DocumentChunk(document_id=document.id, chunk_index=index, page_number=page_number, content=content, embedding=embedding))
                document.status, document.page_count = "ready", unit_count
            except Exception as exc:  # noqa: BLE001
                document.status, document.error = "failed", str(exc)
            session.commit()
            if document.status == "failed":
                raise RuntimeError(document.error or "Không thể index tài liệu.")
            return document

    def delete(self, document_id: str) -> bool:
        with self.database.session() as session:
            document = session.get(Document, document_id)
            if document is None:
                return False
            self.storage.delete(document.storage_provider, document.stored_name, document.storage_file_id)
            session.delete(document)
            session.commit()
            return True

    def search(
        self,
        query: str,
        top_k: int = 4,
        project_id: str | None = None,
        collection_id: str | None = None,
        max_distance: float | None = None,
    ) -> str:
        if not query.strip():
            return "[Lỗi] Câu hỏi truy vấn đang trống."
        top_k = max(1, min(int(top_k), 8))
        vector = self._embed([query], "query")[0]
        with self.database.session() as session:
            distance = DocumentChunk.embedding.cosine_distance(vector)
            statement = (
                select(DocumentChunk, Document.original_name, distance.label("distance"))
                .join(Document)
                .where(Document.status == "ready", Document.project_id == project_id)
            )
            if collection_id:
                statement = statement.join(KnowledgeCollectionDocument).where(KnowledgeCollectionDocument.collection_id == collection_id)
            if max_distance is not None:
                statement = statement.where(distance <= max_distance)
            rows = session.execute(statement.order_by(distance).limit(top_k)).all()
        if not rows:
            return NO_DOCUMENTS_RESULT
        return "\n\n".join(
            (
                f"[Nguồn {number}: [{name}](/api/documents/{chunk.document_id}/file#page={chunk.page_number}), trang {chunk.page_number}]\n{chunk.content}"
                if Path(name).suffix.lower() == ".pdf"
                else f"[Nguồn {number}: [{name}](/api/documents/{chunk.document_id}/file), đoạn {chunk.page_number}]\n{chunk.content}"
            )
            for number, (chunk, name, _) in enumerate(rows, start=1)
        )


def build_knowledge_tool(service: KnowledgeService, project_id: str | None = None, collection_id: str | None = None) -> ToolSpec | None:
    # Chats outside a project use the global document library. Project chats
    # remain opt-in and only search the explicitly selected collection.
    if project_id is not None and not collection_id:
        return None
    return ToolSpec(
        name="search_knowledge_base",
        description="Tìm thông tin trong tài liệu PDF, DOCX hoặc Markdown người dùng đã upload và index. Dùng khi câu hỏi liên quan đến tài liệu; câu trả lời phải nêu nguồn.",
        parameters={"type": "object", "properties": {"query": {"type": "string", "description": "Câu truy vấn rõ ràng bằng ngôn ngữ tự nhiên."}, "top_k": {"type": "integer", "description": "Số đoạn cần lấy, từ 1 đến 8; mặc định 4."}}, "required": ["query"]},
        func=lambda query, top_k=4: service.search(query, top_k, project_id, collection_id),
    )
